# -*- coding: utf-8 -*-
"""Генерація повного розділу 4 + додатки Л, М → DOCX."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUT = Path(__file__).parent / "rozd4_hikora_POVNY.docx"
TEXT_MD = Path(__file__).parent / "rozd4-text-dlya-word.md"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)


def h1(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def h2(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def h3(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def body(t):
    if not t.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(t.strip())
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def code(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(t.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(10)


def cap(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.italic = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def add_table(cap_text, headers, rows):
    cap(cap_text)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(12)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = val
    doc.add_paragraph()


def import_markdown(path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_buf = []
    para_buf = []

    def flush_para():
        nonlocal para_buf
        if para_buf:
            text = " ".join(para_buf)
            if text.strip() and not text.startswith("*Кінець"):
                body(text)
            para_buf = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            code("\n".join(code_buf))
            code_buf = []

    for line in lines:
        raw = line.rstrip()
        s = raw.strip()

        if s.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_para()
                in_code = True
            continue
        if in_code:
            code_buf.append(raw)
            continue

        if not s or s == "---":
            flush_para()
            continue
        if s.startswith("# ") and not s.startswith("## "):
            flush_para()
            h1(s[2:].strip())
            continue
        if s.startswith("## "):
            flush_para()
            h2(s[3:].strip())
            continue
        if s.startswith("### "):
            flush_para()
            h3(s[4:].strip())
            continue
        if s.startswith("|") or s.startswith("*") and "Файл:" in s:
            continue
        if s.startswith("**Таблиця") or s.startswith("**Рис."):
            flush_para()
            cap(s.replace("**", ""))
            continue
        para_buf.append(s)

    flush_para()
    flush_code()


# Пропускаємо дубль заголовка з MD (додамо свій вступ)
h1("РОЗДІЛ 4. РЕАЛІЗАЦІЯ ТА ТЕСТУВАННЯ МОБІЛЬНОГО ЗАСТОСУНКУ HIKORA")
body(
    "У розділі подано програмну реалізацію проєктних рішень (розд. 3), звіт про тестування "
    "та демонстрацію застосування Hikora. Реалізація описана поступово: база даних → "
    "підключення бекенду Supabase → Edge Functions → клієнт Flutter. Тестування виконано "
    "у порядку: функціональне → безпека (ST, MobSF) → навантажувальне (JMeter). "
    "Протоколи — додаток Л; інструкція користувача — додаток М."
)

add_table(
    "Таблиця 4.0. Відповідність функцій специфікації та реалізації (сервер)",
    ["Функція", "Реалізація"],
    [
        ["Автентифікація", "Supabase Auth, RLS"],
        ["Збереження маршруту", "Edge save-route"],
        ["Маршрутизація", "route-hike, GraphHopper/OSRM"],
        ["Рекомендації / ШІ", "recommend-routes, ai-chat"],
        ["Офлайн-трек", "prepare-offline-route"],
        ["Групові походи", "trip-actions, trip-chat"],
        ["Погода, геопошук, POI", "weather, geosearch, poi-nearby"],
        ["Ізоляція даних", "RLS (15 таблиць)"],
    ],
)

add_table(
    "Таблиця 4.0а. Відповідність функцій та клієнтської реалізації",
    ["Функція", "Компонент"],
    [
        ["Вхід / OAuth", "LoginScreen, GoRouter"],
        ["Каталог маршрутів", "RoutesRepository"],
        ["Офлайн / навігація", "OfflineMapService, NavigationScreen"],
        ["Походи / чат", "TripsApi, TripChatScreen"],
        ["ШІ / рекомендації", "AiService, HomeScreen"],
        ["Журнал", "JournalScreen"],
    ],
)

if TEXT_MD.exists():
    import_markdown(TEXT_MD)

doc.add_page_break()

# ─── ДОДАТОК Л ───
h1("ДОДАТОК Л")
h2("Звіт про виконання тестів")
body(
    "Додаток містить протоколи функціонального, безпечного та навантажувального тестування. "
    "До додатка вкладаються скріншоти екранів застосунку, Aggregate Report JMeter, "
    "звіт MobSF (PDF) та вивід flutter analyze."
)

h3("Л.1. Функціональне тестування")
add_table(
    "Таблиця Л.1. Зведені результати (63 випадки)",
    ["Модуль", "Всього", "Пройдено", "Провалено"],
    [
        ["Auth / профіль", "12", "12", "0"],
        ["Маршрути", "14", "14", "0"],
        ["Навігація / офлайн", "11", "11", "0"],
        ["Групові походи", "10", "10", "0"],
        ["Журнал", "8", "8", "0"],
        ["Профіль / досягнення", "8", "8", "0"],
        ["Разом", "63", "63", "0"],
    ],
)

h3("Л.2. Безпека та SAST")
add_table(
    "Таблиця Л.2. Тести безпеки (ST)",
    ["ID", "Опис", "Статус"],
    [
        ["ST-01", "Edge без JWT", "OK"],
        ["ST-03", "Чужий journal", "OK"],
        ["ST-11", "Чат без approved", "OK"],
        ["ST-29", "Немає секретів у lib", "OK"],
        ["ST-30", "JWT у recommend-routes", "OK"],
        ["ST-Storage-1/2", "Storage RLS", "OK"],
        ["Усього", "37 сценаріїв", "37 OK"],
    ],
)
body("MobSF (app-release.apk): HTTPS, no trackers; flutter analyze — без критичних помилок.")

h3("Л.3. Навантаження (JMeter)")
add_table(
    "Таблиця Л.3. JMeter, 1000 VU (орієнтир)",
    ["Метрика", "Значення"],
    [
        ["HTTP-запитів", "13 003"],
        ["Помилок", "13,67 %"],
        ["Піки помилок", "recommend-routes, poi-nearby"],
    ],
)

doc.add_page_break()

# ─── ДОДАТОК М ───
h1("ДОДАТОК М")
h2("Інструкція користувача")

h3("М.1. Компоненти ПЗ")
body("Застосунок Hikora, обліковий запис Supabase, Інтернет, GPS, камера/галерея.")

h3("М.2. Встановлення")
body("Android: APK; iOS: TestFlight. Мінімум Android 10.")

h3("М.3. Налаштування")
body("Реєстрація → профіль → дозволи GPS і фото.")

h3("М.4. Базові функції")
add_table(
    "Таблиця М.1. Дії користувача",
    ["Задача", "Шлях у застосунку"],
    [
        ["Маршрути", "Вкладка «Маршрути»"],
        ["Офлайн", "Деталі → Завантажити; Профіль → Офлайн"],
        ["Походи", "Вкладка «Групи»"],
        ["Журнал", "Після походу або вкладка Журнал"],
    ],
)

h3("М.5. Помилки")
add_table(
    "Таблиця М.2. Типові помилки",
    ["Ситуація", "Дія"],
    [
        ["Немає карти без мережі", "Завантажити офлайн-пакет"],
        ["Немає чату", "Дочекатися схвалення заявки"],
    ],
)

doc.save(OUT)
print(f"Збережено: {OUT}")
