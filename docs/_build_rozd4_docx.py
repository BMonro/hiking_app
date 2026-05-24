"""Build corrected rozdim 4 docx from structured sections."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"d:\Downloads\rozd4_hikora_vypravleno.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

# Code font: 12 pt (main text 14 pt) — readable in print/PDF.
CODE_FONT_PT = 12

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(CODE_FONT_PT)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

# --- Title ---
h1("РОЗДІЛ 4. РЕАЛІЗАЦІЯ ТА ТЕСТУВАННЯ МОБІЛЬНОГО ЗАСТОСУНКУ HIKORA")

# --- 4.1 ---
h2("4.1. Реалізація серверної частини")
body(
    "Серверну частину застосунку розгорнуто на хмарній платформі Supabase без використання "
    "окремого VPS-сервера. База даних PostgreSQL містить 15 таблиць та одне представлення "
    "profile_stats з увімкненим механізмом Row Level Security для розмежування доступу між "
    "користувачами. Таблиці охоплюють усі сутності системи: profiles, profile_health_conditions, "
    "routes, route_points, route_ratings, saved_routes, offline_routes, trips, trip_participants, "
    "messages, journal_entries, journal_photos, achievements, user_achievements та notifications. "
    "Облікові записи зберігаються в auth.users (Supabase Auth)."
)
body(
    "Автентифікація реалізована через Supabase Auth з підтримкою входу через email і пароль, "
    "Google OAuth та відновлення пароля. Після першого входу через Google користувач заповнює "
    "фізичний профіль (вік, рівень підготовки, медичні обмеження)."
)
body(
    "Серверна логіка, що потребує захищеного доступу до зовнішніх API, винесена в Edge Functions "
    "на базі Deno (шар BFF — backend for frontend). Ключі OpenAI, GraphHopper, OpenWeatherMap "
    "та інші зберігаються лише в Supabase Secrets і не потрапляють у клієнтський код. "
    "Розгорнуті функції: ai-chat, recommend-routes, trip-actions, route-hike, "
    "prepare-offline-route, save-route, weather, geosearch, poi-nearby та trip-chat."
)
body(
    "Обмін даними між клієнтом і сервером відбувається через Supabase Dart SDK (PostgREST): "
    "репозиторії формують запити .from('table').select() / .eq() / .upsert(), результати "
    "перетворюються на Dart-моделі вручну в шарі data (патерн Repository, без ORM). "
    "Виклики зовнішніх сервісів з клієнта здійснюються переважно через клас BackendApi "
    "(обгортка над Supabase Edge Functions); бібліотека Dio використовується для завантаження "
    "тайлів OpenStreetMap та як резервний канал при недоступності Edge Functions."
)
body(
    "Для оновлення списку походів використовується Supabase Realtime через WebSocket-підписки "
    "(tripsRealtimeSyncProvider). Завантажені офлайн-маршрути дублюються записом у таблиці "
    "offline_routes, а файли пакета (тайли, route_path.json, map_meta.json, маркер .complete) "
    "зберігаються у файловій системі пристрою в каталозі offline_tiles/{routeId}/ через path_provider."
)
body(
    "Аватари користувачів зберігаються в Supabase Storage (bucket avatars) через сервіс "
    "AvatarStorageService. Обкладинки маршрутів зберігаються як URL у полі cover_image_url таблиці routes."
)
body("Структуру серверної частини наведено на рис. 4.1.")
caption("Рис. 4.1. Структура серверної частини застосунку Hikora")

# --- 4.2 ---
h2("4.2. Реалізація клієнтської частини")
body(
    "Клієнтський застосунок реалізовано мовою Dart у фреймворку Flutter SDK 3.0+. Проєкт "
    "організовано за принципом feature-first: каталог lib/features/ містить модулі auth, home, "
    "routes, navigation, weather, group_hikes, trips, journal, profile та ai. Кожен модуль "
    "поділено на три шари: presentation (екрани та віджети), data (репозиторії та сервіси) "
    "та domain (моделі). Спільна інфраструктура зосереджена в lib/core/ — конфігурація "
    "go_router, тема застосунку, оболонка MainShell з нижньою панеллю навігації та клас BackendApi."
)
body(
    "Управління станом реалізовано бібліотекою flutter_riverpod. Навігація між екранами "
    "здійснюється через go_router із захистом маршрутів. Нижня панель містить вкладки: "
    "Головна, Маршрути, Карта, Групи та Профіль."
)
body("Структуру клієнтської частини подано на рис. 4.2.")
caption("Рис. 4.2. Структура клієнтської частини застосунку")
body("Нижче наведено фрагмент перевірки сесії та перенаправлення неавторизованого користувача.")
caption("Фрагмент перевірки сесії та перенаправлення неавторизованого користувача")
code_block(
    """redirect: (context, state) async {
  final session = Supabase.instance.client.auth.currentSession;
  final loc = state.matchedLocation;
  if (session == null) {
    if (loc == '/login' || loc == '/register') return null;
    return '/login';
  }
  // перевірка заповнення фізичного профілю (age у profiles)
  return null;
},"""
)
body("Нижче наведено фрагмент реалізації входу через Google OAuth.")
caption("Фрагмент екрану авторизації — OAuth через Supabase Auth")
code_block(
    """final launched = await Supabase.instance.client.auth.signInWithOAuth(
  OAuthProvider.google,
  redirectTo: kIsWeb ? null : 'io.supabase.flutter://login-callback/',
);"""
)

h2("4.2.1. Реалізація маршрутів та геопошуку")
body(
    "Каталог публічних маршрутів завантажується через RoutesRepository з фільтрацією за "
    "складністю, типом маршруту (route_type), тривалістю, набором висоти та пошуком за назвою. "
    "Точки маршруту (route_points) завантажуються окремим запитом у методі getRouteDetail. "
    "Збереження нового маршруту виконується через Edge Function save-route, яка обчислює "
    "метрики та записує geojson і точки в БД."
)
body("Нижче наведено фрагмент формування запиту до PostgREST з фільтрами.")
caption("Фрагмент репозиторію маршрутів — формування запиту до PostgREST з фільтрами")
code_block(
    """Future<List<RouteModel>> getRoutes({...}) async {
  dynamic query = _client.from('routes').select().eq('is_public', true);
  if (search != null && search.isNotEmpty) {
    query = query.ilike('title', '%$search%');
  }
  if (difficulty != null && difficulty != 'all') {
    query = query.eq('difficulty', difficulty);
  }
  if (routeType != null && routeType != 'all') {
    query = query.eq('route_type', routeType);
  }
  query = query.order('created_at', ascending: false);
  final data = await query;
  return (data as List).map((json) => RouteModel.fromJson(json)).toList();
}"""
)
body(
    "Пошук назв місць реалізовано через OsmNominatimService, який викликає Edge Function geosearch "
    "(з резервним прямим зверненням до Nominatim та Overpass API для вершин). Побудова пішохідного "
    "маршруту виконується через RoutingRepository та Edge Function route-hike (GraphHopper з "
    "профілем hike, резерв OSRM на сервері); при недоступності Edge — локальний fallback через Dio."
)

h2("4.2.2. Реалізація офлайн-карт та навігації")
body(
    "Сервіс OfflineMapService реалізує метод downloadRouteMap, який повертає потік прогресу "
    "OfflineMapDownloadProgress. У каталозі offline_tiles/{routeId} зберігаються тайли "
    "OpenStreetMap (рівні масштабу 11–16), файл лінії route_path.json, метадані map_meta.json "
    "та маркер .complete. Полілінія для пакета формується через Edge Function prepare-offline-route "
    "(клас PrepareOfflineApi). Після завершення завантаження запис дублюється в таблиці offline_routes."
)
body("Нижче наведено фрагмент завантаження офлайн-пакета.")
caption("Фрагмент сервісу офлайн-карт — завантаження тайлів і збереження лінії маршруту")
code_block(
    """Stream<OfflineMapDownloadProgress> downloadRouteMap(
  RouteDetail detail, {required List<LatLng> pathPolyline},
) async* {
  final routeId = detail.route.id;
  final dir = await _routeDir(routeId); // offline_tiles/{routeId}
  // завантаження тайлів OSM (zoom 11–16) через Dio
  await _writePath(routeId, OfflineRoutePath(...));
  await _writeMeta(routeId, _MapMeta(...));
  await marker.writeAsString('1', flush: true); // .complete
}"""
)
body(
    "Реалізовано два режими навігації: класична онлайн-навігація з екрана деталей маршруту "
    "(перебудова треку через route-hike, шар POI через poi-nearby) та офлайн-only режим із "
    "вкладки Профіль → Мої маршрути → Офлайн з параметром offline=true. На екрані навігації "
    "відображається поточна позиція GPS з обертанням маркера за курсом, пройдена частина "
    "маршруту (сірий колір) та залишкова (зелений). При відхиленні в онлайн-режимі система "
    "показує SnackBar та автоматично перебудовує трек."
)
caption("Рис. 4.3. Екран навігації під час активного походу")

h2("4.2.3. Реалізація ШІ-функціоналу")
body(
    "Модуль штучного інтелекту реалізовано через Edge Functions recommend-routes та ai-chat. "
    "Функція recommend-routes повертає список recommendations (route_id та reason) і поле source "
    "(ai або profile). Клієнтський AiService викликає функцію з порожнім body — автентифікація "
    "за JWT на сервері. Далі personalizedRoutesProvider завантажує картки через getRoutesByIds. "
    "Модель OpenAI gpt-4o-mini; ключ OPENAI_API_KEY — лише в Supabase Secrets. За відсутності "
    "ключа сервер повертає source: profile і список від rankRoutesByProfile."
)
caption("Фрагмент сервісу ШІ — виклик recommend-routes")
code_block(
    """Future<({List<Map<String, String>> items, String source})>
    fetchRecommendations() async {
  final response = await _client.functions.invoke('recommend-routes', body: {});
  final data = _asMap(response.data);
  final source = data?['source']?.toString() ?? 'profile';
  final list = data?['recommendations'];
  return (items: out, source: source);
}"""
)
body(
    "Діалоговий консультант на головному екрані реалізовано функцією ai-chat з історією "
    "повідомлень у форматі JSON."
)

h2("4.2.4. Реалізація групових походів")
body(
    "Створення походу реалізовано через форму з вибором маршруту, дат, опису та ліміту учасників. "
    "Заявки та рішення організатора обробляються Edge Function trip-actions (дії apply, decide, cancel). "
    "Список походів оновлюється через Supabase Realtime. Груповий чат — таблиця messages; доступ "
    "лише для учасників зі статусом approved (RLS). Екран групових походів — GroupHikesScreen, "
    "маршрут /trips, вкладка Групи нижньої панелі."
)

h2("4.2.5. Реалізація погоди, журналу та профілю")
body(
    "Модуль погоди: WeatherRepository викликає Edge Function weather (OpenWeatherMap на сервері); "
    "екрани /weather та погода на маршруті (/routes/detail/:id/weather). Журнал: journal_entries "
    "та journal_photos; після завершення маршруту — діалог збереження запису. Профіль: редагування, "
    "статистика (VIEW profile_stats), досягнення (автонарахування тригерами в БД), налаштування, "
    "вкладка Мої маршрути з офлайн-пакетами."
)

h2("4.3. Результати реалізації застосунку")
body(
    "Після реалізації проведено перевірку на Android-пристрої. Екран авторизації підтримує email "
    "та Google (рис. 4.4). Після реєстрації — налаштування профілю."
)
caption("Рис. 4.4. Екран авторизації застосунку")
body(
    "На головному екрані — рекомендації маршрутів та ШІ-консультант (рис. 4.5). При відсутності "
    "OPENAI_API_KEY джерело рекомендацій — profile."
)
caption("Рис. 4.5. Головний екран з рекомендаціями та ШІ-консультантом")
body("Каталог маршрутів з фільтрами та деталі з офлайн-завантаженням (рис. 4.6).")
caption("Рис. 4.6. Екран деталей маршруту")
body("Групові походи та чат (рис. 4.7).")
caption("Рис. 4.7. Екран групових походів та чату")
body("Журнал після походу (рис. 4.8).")
caption("Рис. 4.8. Екран журналу походів")

h2("4.4. Тестування застосунку")
h2("4.4.1. Розробка тестів")
body(
    "Тестування на Android-емуляторі та фізичному смартфоні. 12 тестових випадків, flutter analyze, "
    "RLS на тестовому Supabase, seed_data.sql. Додаток Л."
)
body("Таблиця 4.1 — Результати функціонального тестування (TC-01 … TC-12 — Пройдено).")
body("Таблиця 4.2 — Результати тестування безпеки (RLS, trip-actions 401, відсутність ключа в клієнті).")
body("Таблиця 4.3 — Орієнтовні показники продуктивності.")

h2("4.4.2. Виявлені та усунені дефекти")
body("Дефект №1. Некоректне закриття підказок OSM — FocusScope.unfocus() після вибору.")
body("Дефект №2. Відсутність обертання маркера GPS — передача heading з geolocator.")
body("Дефект №3. Автоперехід в офлайн — розділено класичну та офлайн-only навігацію.")

h2("4.5. Висновки до розділу")
body(
    "У четвертому розділі описано реалізацію Hikora та результати тестування. Реалізовано "
    "підсистеми: автентифікація, маршрути, офлайн-навігація (два режими), погода, групові походи, "
    "ШІ, журнал, досягнення. Сервер — Supabase (15 таблиць, profile_stats) з Edge Functions BFF."
)
body(
    "Проведено тестування за 12 випадками. У поточній версії немає UI для route_ratings та "
    "saved_routes; FCM не використовується; сповіщення — notifications + Realtime + SnackBar. "
    "Протоколи — додаток Л, інструкція — додаток М."
)

doc.save(OUT)
print(f"Saved: {OUT}")
