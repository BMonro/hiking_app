"""Додатки М і Н — тестові випадки Hikora (кількість за покриттям модуля, не «по 10»)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(r"d:\Downloads\Hikora_dodatok_M_N_v2.docx")
COLS = ("ID", "Назва", "Передумови", "Кроки", "Тестові дані", "Очікуваний результат", "Статус")


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_table(doc, title: str, fig_caption: str, rows: list[tuple], header_fill: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    table = doc.add_table(rows=1 + len(rows), cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, name in enumerate(COLS):
        hdr[i].text = name
        shade_cell(hdr[i], header_fill)
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            if ci == 6:
                shade_cell(cell, "C6EFCE")
            elif ri % 2 == 0:
                shade_cell(cell, "F2F2F2")
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(fig_caption)
    cr.italic = True
    cr.font.size = Pt(12)
    cr.font.name = "Times New Roman"
    summary = doc.add_paragraph()
    sr = summary.add_run(f"Загалом: {len(rows)} тестів | {len(rows)}/{len(rows)}")
    sr.bold = True
    sr.font.size = Pt(10)
    doc.add_paragraph()


def c(name: str, pre: str, steps: str, data: str, expected: str) -> tuple:
    """Один тест без ID — номер присвоїться автоматично."""
    return (name, pre, steps, data, expected)


def numbered_with_next(
    prefix: str, cases: list[tuple], start: int = 1
) -> tuple[list[tuple], int]:
    rows = []
    for i, (name, pre, steps, data, expected) in enumerate(cases, start=start):
        rows.append((f"{prefix}-{i:02d}", name, pre, steps, data, expected, "Успіх ✓"))
    return rows, start + len(cases)


# ─── Додаток М: лише сценарії, що реально є в Hikora ───

M1_AUTH = [
    c("Реєстрація з валідним email і паролем",
      "Застосунок встановлено; email вільний",
      "1. /register → email, пароль ≥6, ім'я.\n2. «Зареєструватись».\n3. Підтвердити email (якщо ввімкнено).",
      "hiker_new@mail.com; Test1234",
      "Акаунт створено; крок 2 профілю або /home"),
    c("Вхід з валідними даними",
      "Обліковий запис існує",
      "1. /login → email і пароль.\n2. «Увійти».",
      "hiker_test@mail.com; Test1234",
      "Сесія JWT; перехід на /home"),
    c("Вхід з невірним паролем",
      "Обліковий запис існує",
      "1. Невірний пароль → «Увійти».",
      "Пароль: WrongPass",
      "Помилка Auth; залишення на /login"),
    c("Вхід через Google OAuth",
      "Google OAuth у Supabase",
      "1. «Увійти через Google».\n2. Обрати акаунт.\n3. Дочекатись deep link callback.",
      "redirectTo: io.supabase.flutter://login-callback/",
      "Сесія створена"),
    c("OAuth без заповненого профілю → /register?oauth=1",
      "Перший вхід Google; age = null",
      "1. Увійти через Google.\n2. Перевірити redirect.",
      "—",
      "Перенаправлення на /register?oauth=1"),
    c("Заповнення кроку 2 (вік, підготовка)",
      "Користувач на кроці 2 реєстрації",
      "1. Вказати age, fitness_level.\n2. Зберегти.",
      "age: 28; fitness_level: intermediate",
      "profiles оновлено; доступ до /home"),
    c("Неспівпадіння паролів при реєстрації",
      "Крок 1 реєстрації",
      "1. Різні пароль і підтвердження.\n2. Продовжити.",
      "password ≠ confirm",
      "FormValidators.confirmPassword: «Паролі не співпадають»"),
    c("Вихід із системи",
      "Авторизований користувач",
      "1. Налаштування → «Вийти».",
      "—",
      "signedOut; redirect /login"),
    c("Захист /home без сесії",
      "Немає JWT",
      "1. Відкрити /home.",
      "—",
      "go_router → /login"),
    c("Відновлення пароля",
      "SMTP Supabase налаштовано",
      "1. «Забули пароль?» → email → надіслати.",
      "hiker_test@mail.com",
      "Лист recovery надіслано"),
]

M2_ROUTES = [
    c("Пошук за назвою",
      "Є публічні маршрути",
      "1. Маршрути → пошук «Карпати».",
      "title ILIKE",
      "Відфільтрований список"),
    c("Фільтр складності та типу разом",
      "Різні difficulty і route_type",
      "1. hard + circular.",
      "difficulty; route_type",
      "Обидва .eq застосовані"),
    c("Фільтр тривалості й набору висоти",
      "Різні duration_h, ascent_m",
      "1. max 5 год, 800 м.",
      "lte duration_h; lte ascent_m",
      "Список відповідає порогам"),
    c("Створення маршруту (старт + фініш)",
      "Є мережа",
      "1. Редактор → точки start/finish.\n2. Зберегти.",
      "save-route create",
      "route_id; distance_km, ascent_m, duration_h, geojson"),
    c("Блокування без фінішу",
      "Редактор відкритий",
      "1. Лише start → «Зберегти».",
      "1 точка start",
      "start_finish_required"),
    c("Геопошук точки (geosearch)",
      "Є інтернет",
      "1. Ввести «Говерла» → обрати підказку.",
      "geosearch",
      "Координати підставлені"),
    c("Редагування власного маршруту",
      "Користувач — author",
      "1. Відкрити свій маршрут → змінити title.\n2. save-route update.",
      "action: update",
      "Маршрут оновлено в БД"),
    c("Видалення власного маршруту",
      "Користувач — author",
      "1. Мої маршрути / деталі → видалити.\n2. Підтвердити.",
      "deleteRoute",
      "Маршрут і route_points видалені (CASCADE)"),
    c("Приватний маршрут невидимий іншим",
      "is_public = false",
      "1. Створити приватний.\n2. Увійти іншим акаунтом.",
      "—",
      "Не в публічному каталозі (RLS)"),
    c("Деталі: карта, точки, метрики",
      "Маршрут з geojson",
      "1. Відкрити картку.",
      "—",
      "Трек, waypoints, distance/ascent/duration"),
    c("Погода на маршруті",
      "Edge weather доступна",
      "1. Деталі → «Погода».",
      "action: both",
      "Поточна + прогноз"),
    c("Екран /weather за містом",
      "Авторизований",
      "1. Вкладка погоди / маршрут weather.\n2. Пошук міста.",
      "action: city",
      "WeatherModel відображено"),
    c("Побудова треку (route-hike) при збереженні",
      "≥2 waypoints",
      "1. Зберегти маршрут.\n2. Перевірити geojson у БД.",
      "fetchHikingRouteThrough",
      "LineString ≥2 координат"),
]

M3_NAV = [
    c("Завантаження офлайн-пакета",
      "Деталі маршруту; Wi-Fi",
      "1. «Завантажити карту офлайн» → 100%.",
      "zoom 11–16",
      ".complete, offline_routes, route_path.json"),
    c("Видалення офлайн-пакета",
      "Пакет завантажено",
      "1. Мої маршрути → видалити офлайн.",
      "deleteOfflineMap",
      "Каталог offline_tiles очищено"),
    c("Офлайн-навігація без мережі",
      "Пакет є; мережа вимкнена",
      "1. Офлайн → запуск навігації.",
      "offline=true",
      "Локальні тайли; GPS працює"),
    c("Онлайн-навігація та GPS heading",
      "Дозвіл геолокації",
      "1. «Почати проходження».\n2. Рух пристроєм.",
      "—",
      "Маркер рухається; обертання за heading"),
    c("Reroute при відхиленні від треку",
      "Онлайн; route-hike доступний",
      "1. Відійти >75 м від сегмента.",
      "пороги 45/75 м",
      "SnackBar; новий polyline"),
    c("Пройдений і залишковий трек",
      "Навігація активна",
      "1. Пройти частину шляху.",
      "—",
      "Сірий / зелений колір ліній"),
    c("POI на карті (онлайн)",
      "poi-nearby",
      "1. Увімкнути POI.\n2. Змістити карту.",
      "Overpass через Edge",
      "Маркери POI"),
    c("Діалог завершення → журнал",
      "Навігація завершена",
      "1. «Завершити» → зберегти в журнал.",
      "actual_* метрики",
      "Діалог з prefilled метриками"),
    c("Два режими: онлайн vs offline-only",
      "Є пакет + мережа",
      "1. Старт з деталей.\n2. Старт з вкладки Офлайн.",
      "—",
      "Онлайн: reroute; офлайн: без HTTP"),
]

M4_TRIPS = [
    c("Створення групового походу",
      "Авторизований",
      "1. Групи → створити → форма.\n2. Зберегти.",
      "trip-actions create",
      "trip_id, trip_code; status open"),
    c("Редагування походу організатором",
      "Організатор; похід open",
      "1. Редагувати title/dates.\n2. update.",
      "action: update",
      "trips оновлено"),
    c("Подача заявки",
      "Не організатор; open",
      "1. «Подати заявку».",
      "apply",
      "pending; notification організатору"),
    c("Схвалення заявки",
      "pending існує",
      "1. «Схвалити».",
      "decide approved",
      "approved; доступ до чату"),
    c("Відхилення заявки",
      "pending існує",
      "1. «Відхилити».",
      "decide rejected",
      "rejected; сповіщення учаснику"),
    c("Організатор не подає заявку на свій похід",
      "organizer_id = self",
      "1. Спроба apply.",
      "—",
      "organizer_cannot_apply"),
    c("Похід заповнений (max_members)",
      "approved = max_members",
      "1. Нова заявка.",
      "trip_full",
      "Помилка trip_full"),
    c("Повідомлення в чаті",
      "approved",
      "1. Надіслати текст.",
      "trip-chat / messages",
      "Повідомлення в UI + Realtime"),
    c("Чат недоступний для pending",
      "status pending",
      "1. Відкрити чат / SELECT messages.",
      "—",
      "RLS блокує"),
    c("Realtime оновлення списку походів",
      "2 клієнти",
      "1. Створити похід на A.\n2. Переглянути список на B.",
      "tripsRealtimeSyncProvider",
      "Список оновився без restart"),
    c("Скасування походу",
      "Організатор",
      "1. cancel.",
      "action: cancel",
      "status cancelled"),
    c("Профіль заявника для організатора",
      "pending заявка",
      "1. Торкнутись імені заявника.",
      "public_profile",
      "Превʼю профілю (age, fitness, bio)"),
]

M5_AI = [
    c("Рекомендації на головній",
      "Є публічні маршрути",
      "1. /home → блок рекомендацій.",
      "recommend-routes",
      "До 5 карток; source ai|profile"),
    c("Перехід у деталі з рекомендації",
      "Є recommendations",
      "1. Торкнутись картки.",
      "getRoutesByIds",
      "Route details відкрито"),
    c("ШІ-чат: запит і відповідь",
      "OPENAI_API_KEY є",
      "1. Питання в ai-chat → надіслати.",
      "ai-chat",
      "reply у стрічці"),
    c("Резерв без OpenAI",
      "Ключ відсутній",
      "1. Відкрити /home.",
      "—",
      "source: profile; rankRoutesByProfile"),
    c("Недоступний ai-chat",
      "Edge не розгорнуто",
      "1. Надіслати повідомлення.",
      "—",
      "Повідомлення про недоступність; без crash"),
]

M6_JOURNAL = [
    c("Запис після навігації",
      "Діалог завершення",
      "1. Зберегти з метриками.",
      "journal_entries",
      "Запис створено"),
    c("Ручне створення запису",
      "Журнал відкритий",
      "1. «+» → дата, title, notes → зберегти.",
      "—",
      "Новий рядок journal_entries"),
    c("До 5 фото в записі",
      "Редагування запису",
      "1. Додати 5 фото.\n2. Спробувати 6-те.",
      "max 5 images",
      "Кнопка додавання заблокована на 6-му"),
    c("Видалення запису журналу",
      "Є запис",
      "1. Видалити → підтвердити.",
      "—",
      "Запис і journal_photos видалені"),
    c("Редагування notes",
      "Запис існує",
      "1. Редагувати → змінити notes.",
      "—",
      "Оновлено в БД"),
    c("Статистика profile_stats",
      "≥1 похід у журналі",
      "1. Профіль → Статистика.",
      "VIEW profile_stats",
      "Суми km/ascent/hikes; графік"),
    c("Досягнення після порогу",
      "Умова hikes_count",
      "1. Профіль → Досягнення.",
      "first_hike",
      "user_achievements містить badge"),
]

M7_PROFILE = [
    c("In-app сповіщення trip_request",
      "Організатор онлайн",
      "1. Подати заявку з іншого акаунту.",
      "notifications Realtime",
      "SnackBar / список /notifications"),
    c("Позначити сповіщення прочитаним",
      "Є непрочитані",
      "1. /notifications → відкрити.\n2. markRead.",
      "is_read: true",
      "Бейдж зменшився"),
    c("Завантаження аватара",
      "bucket avatars",
      "1. Редагувати профіль → фото.",
      "{userId}-avatar.jpg",
      "avatar_url оновлено"),
    c("Медичні обмеження (1НФ)",
      "Профіль",
      "1. Додати 2 condition.\n2. Зберегти.",
      "profile_health_conditions",
      "2 рядки; UNIQUE user+condition"),
    c("Публічний профіль учасника походу",
      "approved у trip",
      "1. Превʼю з чату.",
      "profiles_public_read",
      "full_name, age, fitness (без секретів)"),
    c("Мої маршрути: авторські та офлайн",
      "Є власний + офлайн маршрут",
      "1. Профіль → Мої маршрути.\n2. Вкладки.",
      "—",
      "Списки розділені; офлайн з позначкою"),
    c("Налаштування сповіщень",
      "Settings",
      "1. Вимкнути тип сповіщень.\n2. Подія → перевірити SnackBar.",
      "notification_preferences",
      "Тип приглушено згідно prefs"),
]

# ─── Додаток Н: групи за типом перевірки, не за кількістю ───

N1_ACCESS = [
    c("Edge Function без JWT",
      "—", "1. POST trip-actions без Authorization.", "—", "401 unauthorized"),
    c("Невалідний JWT на route-hike",
      "—", "1. invoke з Bearer fake.", "—", "401"),
    c("Чужий journal_entries (RLS)",
      "JWT A", "1. SELECT journal B.", "user_id B", "[]"),
    c("messages без approved",
      "pending", "1. SELECT messages.", "—", "RLS блок"),
    c("PATCH чужого routes",
      "author B", "1. PATCH routes id B.", "JWT A", "0 rows"),
    c("offline_routes чужого user",
      "—", "1. SELECT з фільтром B.", "JWT A", "лише свої"),
    c("notifications чужого user",
      "—", "1. SELECT notifications.", "JWT A", "лише свої"),
    c("Подвійна заявка apply",
      "вже pending", "1. apply знову.", "—", "already_applied"),
    c("organizer_cannot_apply",
      "organizer", "1. apply на свій trip.", "—", "помилка Edge"),
    c("decide не організатором",
      "JWT учасника", "1. decide чужої заявки.", "—", "forbidden"),
    c("INSERT messages не учасником",
      "не в trip", "1. INSERT messages.", "—", "RLS / Edge блок"),
    c("Приватний route SELECT",
      "is_public false", "1. SELECT як не-автор.", "—", "немає в результаті"),
    c("save-route без author на update",
      "route B", "1. update чужого route_id.", "JWT A", "403 forbidden"),
]

N2_VALIDATION = [
    c("Некоректний email",
      "register", "1. not-an-email.", "—", "FormValidators.email"),
    c("Пароль <6",
      "register", "1. 5 символів.", "Abc12", "Мінімум 6"),
    c("Паролі не співпадають",
      "register", "1. різні confirm.", "—", "confirmPassword"),
    c("Порожній title походу",
      "форма trip", "1. title пустий.", "—", "не відправлено"),
    c("Назва маршруту >200",
      "редактор", "1. 201 символ.", "—", "title max 200"),
    c("Вік поза 1–119",
      "профіль", "1. age 0 або 200.", "—", "помилка age"),
    c("save-route без start/finish",
      "JWT", "1. POST лише viewpoint.", "—", "start_finish_required"),
    c("API trip title 1 символ",
      "JWT", "1. create title «А».", "—", "400"),
    c("max_members від'ємний",
      "JWT", "1. create max_members -1.", "—", "відхилено"),
    c("XSS у description trip",
      "форма", "1. <script>…\n2. Перегляд.", "—", "екранування"),
    c("Порожній body save-route",
      "JWT", "1. POST {}.", "—", "title_required"),
]

N3_ATTACKS = [
    c("SQLi в пошуку routes",
      "auth", "1. пошук ' OR 1=1 --.", "—", "параметризований запит"),
    c("Stored XSS у journal notes",
      "—", "1. зберегти <script>.\n2. відкрити.", "—", "текст, не скрипт"),
    c("Brute-force login",
      "—", "1. 10+ невірних паролів.", "—", "rate limit Auth"),
    c("IDOR trip_id",
      "JWT", "1. перебір UUID trips.", "—", "немає чужих приватних"),
    c("OPENAI_KEY у APK",
      "release", "1. пошук sk- у lib.", "—", "відсутній"),
    c("Підміна user_id у recommend body",
      "JWT A", "1. body user_id B.", "—", "профіль з JWT A"),
    c("DELETE чужих journal_photos",
      "JWT A", "1. DELETE photo B.", "—", "RLS 0"),
]

N4_FILES = [
    c("Аватар не-зображення",
      "профіль", "1. .exe як фото.", "—", "відхилено"),
    c("Фото журналу >ліміту",
      "журнал", "1. великий файл.", "12 MB", "помилка upload"),
    c("Порожній файл",
      "журнал", "1. 0 bytes.", "—", "не збережено"),
    c("Storage avatars без JWT",
      "—", "1. upload без сесії.", "—", "403"),
    c("Шлях avatars чужого uid",
      "JWT A", "1. B-avatar.jpg.", "—", "RLS uid-%"),
    c("6-те фото журналу (клієнт)",
      "редактор", "1. 5 фото OK.\n2. 6-те.", "—", "кнопка disabled"),
]


def main() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(14)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("Додаток М. Тестові випадки для функціонального тестування системи Hikora")
    run.bold = True
    run.font.size = Pt(14)

    ft_start = 1
    sections_m = [
        ("FT — Автентифікація та профіль. Функціональні тести — модуль «Автентифікація»",
         "Рис. М.1. Тестові випадки для тестування автентифікації та авторизації",
         M1_AUTH, "2E5090"),
        ("FT — Маршрути. Функціональні тести — модуль «Маршрути»",
         "Рис. М.2. Тестові випадки для тестування створення та роботи з маршрутами",
         M2_ROUTES, "2E7D32"),
        ("FT — Офлайн-карти та навігація. Функціональні тести — модуль «Навігація»",
         "Рис. М.3. Тестові випадки для тестування офлайн-пакетів та GPS-навігації",
         M3_NAV, "00695C"),
        ("FT — Групові походи та чат. Функціональні тести — модуль «Групові походи»",
         "Рис. М.4. Тестові випадки для тестування групових походів",
         M4_TRIPS, "5D4037"),
        ("FT — ШІ. Функціональні тести — модуль «ШІ-рекомендації»",
         "Рис. М.5. Тестові випадки для тестування ШІ-функціоналу",
         M5_AI, "6A1B9A"),
        ("FT — Журнал. Функціональні тести — модуль «Журнал і статистика»",
         "Рис. М.6. Тестові випадки для тестування журналу походів",
         M6_JOURNAL, "455A64"),
        ("FT — Профіль. Функціональні тести — модуль «Профіль і сповіщення»",
         "Рис. М.7. Тестові випадки для тестування профілю та верифікації даних",
         M7_PROFILE, "1565C0"),
    ]

    for title, cap, cases, color in sections_m:
        rows, ft_start = numbered_with_next("FT", cases, ft_start)
        add_table(doc, title, cap, rows, color)

    doc.add_page_break()

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = h2.add_run("Додаток Н. Тестові випадки для тестування безпеки системи Hikora")
    r2.bold = True
    r2.font.size = Pt(14)

    st_start = 1
    sections_n = [
        ("ST — Доступ. Тести безпеки — розмежування та заборона доступу",
         "Рис. Н.1. Тестові випадки для перевірки механізмів розмежування та заборони доступу",
         N1_ACCESS, "C62828"),
        ("ST — Валідація. Тести безпеки — коректність введених даних",
         "Рис. Н.2. Тестові випадки для перевірки коректності валідації введених користувачем даних",
         N2_VALIDATION, "E65100"),
        ("ST — Атаки. Тести безпеки — типові вектори атак",
         "Рис. Н.3. Тестові випадки для перевірки захисту від типових векторів атак та вразливостей",
         N3_ATTACKS, "AD1457"),
        ("ST — Файли. Тести безпеки — завантаження файлів",
         "Рис. Н.4. Тестові випадки для перевірки механізмів завантаження файлів",
         N4_FILES, "4527A0"),
    ]

    for title, cap, cases, color in sections_n:
        rows, st_start = numbered_with_next("ST", cases, st_start)
        add_table(doc, title, cap, rows, color)

    doc.save(OUT)
    total_ft = ft_start - 1
    total_st = st_start - 1
    print(f"Saved: {OUT}")
    print("Functional by table:", [len(x[2]) for x in sections_m], f"= {total_ft}")
    print("Security by table:  ", [len(x[2]) for x in sections_n], f"= {total_st}")


if __name__ == "__main__":
    main()
