"""Word: великі фрагменти коду з ключовими обчисленнями (розділ 4)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"d:\Downloads\rozd4_fragmenty_kodu.docx")
CODE_PT = 11  # трохи менший кегль, щоб великі блоки вміщались на сторінку
CAPTION_PT = 12

# (заголовок, файл, код)
FRAGMENTS: list[tuple[str, str, str]] = [
    (
        "1. Обчислення метрик маршруту (відстань, набір висоти, тривалість)",
        "supabase/functions/_shared/routing.ts — computeRouteStats, haversineM",
        """function haversineM(a: LatLng, b: LatLng): number {
  const R = 6371000;
  const dLat = ((b.lat - a.lat) * Math.PI) / 180;
  const dLon = ((b.lon - a.lon) * Math.PI) / 180;
  const lat1 = (a.lat * Math.PI) / 180;
  const lat2 = (b.lat * Math.PI) / 180;
  const x =
    Math.sin(dLat / 2) ** 2 +
    Math.cos(lat1) * Math.cos(lat2) * Math.sin(dLon / 2) ** 2;
  return 2 * R * Math.asin(Math.sqrt(x));
}

export function computeRouteStats(
  points: { lat: number; lon: number; altitude_m?: number | null }[],
): { distance_km: number; ascent_m: number; duration_h: number } {
  const coords = points.filter((p) =>
    Number.isFinite(p.lat) && Number.isFinite(p.lon)
  );
  let meters = 0;
  let ascent = 0;
  for (let i = 1; i < coords.length; i++) {
    const prev = coords[i - 1];
    const cur = coords[i];
    meters += haversineM(
      { lat: prev.lat, lon: prev.lon },
      { lat: cur.lat, lon: cur.lon },
    );
    const pa = prev.altitude_m;
    const ca = cur.altitude_m;
    if (pa != null && ca != null) {
      const diff = ca - pa;
      if (diff > 0) ascent += diff;
    }
  }
  const km = meters / 1000;
  const elevationHours = ascent / 600;
  const durationH = km === 0 ? 0 : km / 4 + elevationHours;
  return {
    distance_km: Math.round(km * 100) / 100,
    ascent_m: ascent,
    duration_h: Math.round(durationH * 10) / 10,
  };
}""",
    ),
    (
        "2. Збереження маршруту: метрики + побудова треку (GraphHopper/OSRM)",
        "supabase/functions/save-route/index.ts",
        """const statsInput = points.map((p) => ({
  lat: Number(p.latitude),
  lon: Number(p.longitude),
  altitude_m: p.altitude_m != null ? Number(p.altitude_m) : null,
}));
const stats = computeRouteStats(statsInput);

const waypoints: LatLng[] = statsInput.map((p) => ({ lat: p.lat, lon: p.lon }));
let geojson: Record<string, unknown> | null = null;
try {
  const { points: poly } = await fetchHikingRouteThrough(waypoints);
  if (poly.length >= 2) {
    geojson = toGeoJsonLineString(poly);
  }
} catch (_) {
  if (waypoints.length >= 2) {
    geojson = toGeoJsonLineString(waypoints);
  }
}

const routePayload = {
  title,
  route_type: body.route_type ?? "linear",
  description: body.description?.trim() ?? "",
  difficulty: body.difficulty ?? "easy",
  distance_km: stats.distance_km,
  duration_h: stats.duration_h,
  ascent_m: stats.ascent_m,
  geojson,
};""",
    ),
    (
        "3. Пішохідна маршрутизація: GraphHopper → OSRM → ланцюг сегментів",
        "supabase/functions/_shared/routing.ts — fetchHikingRouteThrough",
        """export async function fetchHikingRouteThrough(
  waypoints: LatLng[],
): Promise<{ points: LatLng[]; source: string }> {
  if (waypoints.length < 2) {
    return { points: [], source: "none" };
  }

  const key = ghKey();
  if (key) {
    try {
      const pts = await graphHopperRoute(waypoints);
      return { points: pts, source: "graphhopper" };
    } catch (_) { /* OSRM */ }
  }

  try {
    const pts = await osrmFootRoute(waypoints);
    return { points: pts, source: "osrm" };
  } catch (_) { /* chain */ }

  const pts = await chainLegRoutes(waypoints);
  return { points: pts, source: key ? "chain" : "osrm_chain" };
}

async function chainLegRoutes(waypoints: LatLng[]): Promise<LatLng[]> {
  const out: LatLng[] = [];
  for (let i = 0; i < waypoints.length - 1; i++) {
    const a = waypoints[i];
    const b = waypoints[i + 1];
    let seg: LatLng[] = [a, b];
    let gotTrail = false;
    if (key) {
      try {
        const g = await graphHopperRoute([a, b]);
        if (g.length >= 2) { seg = g; gotTrail = true; }
      } catch (_) {}
    }
    if (!gotTrail) {
      try {
        const o = await osrmFootRoute([a, b]);
        if (o.length >= 2) { seg = o; gotTrail = true; }
      } catch (_) {}
    }
    appendSegment(out, seg);
  }
  return out;
}""",
    ),
    (
        "4. Резервне ранжування маршрутів за профілем (без OpenAI)",
        "supabase/functions/_shared/route_ranking.ts — rankRoutesByProfile",
        """export function rankRoutesByProfile(
  profile: ProfileRow | null,
  routes: RouteRow[],
  limit = 5,
): Array<{ route_id: string; reason: string }> {
  const fitness = profile?.fitness_level ?? "beginner";
  const prefDiff = profile?.preferred_difficulty;
  const prefDuration = profile?.preferred_duration_h;
  const experience = profile?.experience_count ?? 0;
  const targetDiff = prefDiff ?? difficultyFromFitness(fitness);

  const scored: Array<{ route: RouteRow; score: number; reason: string }> = [];

  for (const route of routes) {
    let score = 0;
    const parts: string[] = [];

    const diffScore = difficultyScore(targetDiff, route.difficulty);
    score += diffScore * 40;
    if (diffScore > 0.7) parts.push("відповідає вашій складності");

    if (prefDuration && prefDuration > 0 && route.duration_h > 0) {
      const ratio = route.duration_h / prefDuration;
      if (ratio <= 1.15) {
        score += 25;
        parts.push("комфортна тривалість");
      } else if (ratio <= 1.4) {
        score += 10;
      } else {
        score -= 15;
      }
    }

    if (experience < 3 && route.difficulty === "easy") {
      score += 15;
      parts.push("підходить для набору досвіду");
    } else if (experience >= 10 && route.difficulty === "hard") {
      score += 10;
      parts.push("цікавий виклик для досвідченого");
    }

    if (route.ascent_m > 0 && route.ascent_m <= 600 && fitness === "beginner") {
      score += 8;
    }

    scored.push({ route, score, reason: parts.join(", ") || "..." });
  }

  scored.sort((a, b) => b.score - a.score);
  return scored.slice(0, limit).map((e) => ({
    route_id: e.route.id,
    reason: e.reason,
  }));
}""",
    ),
    (
        "5. Підготовка полілінії для офлайн-пакета",
        "supabase/functions/prepare-offline-route/index.ts",
        """const waypoints: LatLng[] = (pointRows ?? []).map((p) => ({
  lat: Number(p.latitude),
  lon: Number(p.longitude),
})).filter((p) => Number.isFinite(p.lat) && Number.isFinite(p.lon));

let polyline: LatLng[] = [];
let source = "waypoints";

try {
  const routed = await fetchHikingRouteThrough(waypoints);
  if (routed.points.length >= 2) {
    polyline = routed.points;
    source = routed.source;
  }
} catch (_) { /* fallback */ }

if (polyline.length < 2 && route.geojson) {
  const coords = extractGeoJsonCoords(route.geojson);
  if (coords.length >= 2) {
    polyline = coords;
    source = "geojson";
  }
}

if (polyline.length < 2) {
  polyline = waypoints;
  source = "waypoints";
}

return jsonResponse({
  ok: true,
  source,
  polyline: polyline.map((p) => ({ lat: p.lat, lon: p.lon })),
  waypoints: waypointsOut,
});""",
    ),
    (
        "6. Офлайн-пакет: bbox, перелік тайлів OSM (zoom 11–16), завантаження",
        "lib/features/navigation/data/offline_map_service.dart",
        """Stream<OfflineMapDownloadProgress> downloadRouteMap(
  RouteDetail detail, {
  required List<LatLng> pathPolyline,
}) async* {
  final routeId = detail.route.id;
  final points = [...pathPolyline, ...detail.waypoints.map((w) => w.position)];
  final bounds = _boundsForPoints(points);
  final tiles = <({int z, int x, int y})>[];
  for (var z = minZoom; z <= maxZoom; z++) {
    tiles.addAll(_tilesForBounds(bounds, z));
  }

  var completed = 0;
  yield OfflineMapDownloadProgress(completed: 0, total: tiles.length);

  const batchSize = 6;
  for (var i = 0; i < tiles.length; i += batchSize) {
    final batch = tiles.skip(i).take(batchSize).toList();
    await Future.wait(batch.map((tile) async {
      final file = File('${dir.path}/${tile.z}/${tile.x}/${tile.y}.png');
      if (!await file.exists()) {
        final url = osmTileUrlTemplate
            .replaceAll('{z}', '${tile.z}')
            .replaceAll('{x}', '${tile.x}')
            .replaceAll('{y}', '${tile.y}');
        final bytes = await _dio.get<List<int>>(url,
            options: Options(responseType: ResponseType.bytes));
        await file.parent.create(recursive: true);
        await file.writeAsBytes(bytes.data ?? []);
      }
      completed++;
    }));
    yield OfflineMapDownloadProgress(completed: completed, total: tiles.length);
  }
  await _writePath(routeId, OfflineRoutePath(...));
  await marker.writeAsString('1', flush: true); // .complete
}

List<({int z, int x, int y})> _tilesForBounds(bounds, int zoom) {
  final minX = _lonToTileX(bounds.minLon, zoom);
  final maxX = _lonToTileX(bounds.maxLon, zoom);
  final minY = _latToTileY(bounds.maxLat, zoom);
  final maxY = _latToTileY(bounds.minLat, zoom);
  final tiles = <({int z, int x, int y})>[];
  for (var x = minX; x <= maxX; x++) {
    for (var y = minY; y <= maxY; y++) {
      tiles.add((z: zoom, x: x, y: y));
    }
  }
  return tiles;
}""",
    ),
    (
        "7. Навігація: відстань до треку та автоматична перебудова (reroute)",
        "lib/features/navigation/presentation/navigation_screen.dart",
        """static const double _offRouteThresholdM = 45;
static const double _offRouteRerouteImmediatelyM = 75;

double _distanceToRouteAhead(LatLng? user, List<LatLng> route, int minIdx) {
  if (user == null || route.length < 2) return 0;
  final start = minIdx.clamp(0, route.length - 2);
  var minD = double.infinity;
  for (var i = start; i < route.length - 1; i++) {
    minD = math.min(
      minD,
      _pointToSegmentMeters(user, route[i], route[i + 1]),
    );
  }
  return minD.isFinite ? minD : 0;
}

double _pointToSegmentMeters(LatLng p, LatLng a, LatLng b) {
  const latScale = 111320.0;
  final lonScale =
      111320.0 * math.cos((a.latitude + b.latitude) * math.pi / 360);
  // проєкція точки p на відрізок a–b у метричній площині
  var t = ((px - ax) * dx + (py - ay) * dy) / len2;
  t = t.clamp(0.0, 1.0);
  final closest = LatLng((ay + t * dy) / latScale, (ax + t * dx) / lonScale);
  return dist.as(LengthUnit.Meter, p, closest);
}

Future<void> _maybeRerouteFromPosition(LatLng user) async {
  if (_offlineOnlyNav) return;
  final offM = _distanceToRouteAhead(user, pts, _routeProgressIndex);
  if (offM < _offRouteThresholdM) return;

  final shouldReroute = offM >= _offRouteRerouteImmediatelyM ||
      _offRouteStreak >= 2;
  if (!shouldReroute) return;

  var targets = _remainingNavWaypoints(pts);
  final newRoute = await _routingRepo.fetchHikingRouteThrough([
    user,
    ...targets,
  ]);
  setState(() {
    _routePoints = newRoute;
    _routeProgressIndex = 0;
    _lastRerouteAt = now;
  });
}""",
    ),
    (
        "8. ШІ-рекомендації: OpenAI gpt-4o-mini або резерв rankRoutesByProfile",
        "supabase/functions/recommend-routes/index.ts",
        """const { data: routes } = await supabase
  .from("routes")
  .select("id, title, difficulty, distance_km, duration_h, ascent_m, ...")
  .eq("is_public", true)
  .limit(40);

const { data: profile } = await supabase
  .from("profiles")
  .select("fitness_level, preferred_difficulty, preferred_duration_h, experience_count")
  .eq("id", userId)
  .maybeSingle();

if (tryAi) {
  const profileContext = await buildProfileContext(supabase, userId);
  const raw = await chatCompletion(
    [
      { role: "system", content: system },
      { role: "user", content: `Профіль:\\n${profileContext}\\n\\nКаталог:\\n${JSON.stringify(catalog)}` },
    ],
    { jsonMode: true },
  );
  const parsed = JSON.parse(raw);
  const recommendations = (parsed.recommendations ?? [])
    .filter((item) => item.route_id && item.reason)
    .slice(0, 5);
  if (recommendations.length > 0) {
    return jsonResponse({ recommendations, source: "ai" });
  }
}

const recommendations = rankRoutesByProfile(profile, routeRows, 5);
return jsonResponse({ recommendations, source: "profile" });""",
    ),
]


def main() -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Hikora — ключові обчислення (розділ 4)\n"
        "Великі фрагменти коду для вставки в диплом"
    )
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

    note = doc.add_paragraph()
    nr = note.add_run(
        f"Шрифт коду: Consolas {CODE_PT} pt. Скопіюйте потрібний блок у Word (Ctrl+C → Ctrl+V)."
    )
    nr.font.size = Pt(12)
    nr.italic = True
    nr.font.name = "Times New Roman"

    doc.add_paragraph()

    for caption, source, code in FRAGMENTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        r = p.add_run(f"{caption}\n{source}")
        r.bold = True
        r.font.size = Pt(CAPTION_PT)
        r.font.name = "Times New Roman"

        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Cm(0.4)
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(12)
        cr = cp.add_run(code)
        cr.font.name = "Consolas"
        cr.font.size = Pt(CODE_PT)

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Blocks: {len(FRAGMENTS)}, font: Consolas {CODE_PT} pt")


if __name__ == "__main__":
    main()
