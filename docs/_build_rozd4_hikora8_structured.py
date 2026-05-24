"""
Розділ 4 за вимогами методички (п. 3.12): реалізація проєктних рішень (опис функцій),
демонстрація, тестування (розробка тестів + звіт), посилання на додатки М, Н, Л.
"""
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUT = Path(r"d:\Downloads\rozd4_hikora (8)_структура.docx")
V8_TAIL = Path(__file__).parent / "_rozd4_v8_text.txt"
CODE_PT = 11

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)


def h1(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def h2(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def h3(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def body(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def code(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(CODE_PT)


def caption(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"


def add_table(cap: str, headers: list[str], rows: list[list[str]]) -> None:
    caption(cap)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
    doc.add_paragraph()


# ─── РОЗДІЛ 4 ───
h1("РОЗДІЛ 4. РЕАЛІЗАЦІЯ ТА ТЕСТУВАННЯ МОБІЛЬНОГО ЗАСТОСУНКУ HIKORA")

body(
    "У цьому розділі подано програмну реалізацію проєктних рішень, розроблених у розділі 3, "
    "звіт про тестування та приклади застосування мобільного застосунку Hikora. "
    "Опис зосереджено на реалізованих функціях предметної області: для кожної функції "
    "наведено призначення, програмні компоненти реалізації та ілюстративні уривки коду; "
    "повні тексти програм винесено в додатки та каталог supabase/functions/ відповідно до "
    "узгодження з керівником (п. 3.12.2 методичних рекомендацій)."
)
body(
    "Структура розділу: п. 4.1–4.2 — опис реалізації серверних і клієнтських функцій; "
    "п. 4.3 — демонстрація роботи (п. 3.12.6); п. 4.4 — розробка тестів і звіт про "
    "виконання тестів (п. 3.12.3), з деталізацією у додатках М і Н; п. 4.5 — висновки. "
    "Інструкція користувача подана в додатку Л (п. 3.12.8)."
)

# ═══════════════════════════════════════════════════════════════
# 4.1 СЕРВЕР — реалізовані функції
# ═══════════════════════════════════════════════════════════════
h2("4.1. Реалізація серверної частини системи")
body(
    "Серверну частину реалізовано на платформі Supabase [12]: PostgreSQL [13] "
    "(15 таблиць, представлення profile_stats), Supabase Auth, Storage, Realtime та "
    "10 Edge Functions (Deno) — програмних модулів серверної бізнес-логіки. "
    "Відповідність функцій специфікації (розд. 2) і проєктним рішенням (розд. 3) "
    "наведено в табл. 4.0."
)
add_table(
    "Таблиця 4.0. Відповідність функцій специфікації та програмної реалізації (сервер)",
    ["Функція системи", "Програмна реалізація"],
    [
        ["Реєстрація, вхід, OAuth", "Supabase Auth; політики RLS на profiles"],
        ["Збереження маршруту з метриками та треком", "Edge Function save-route"],
        ["Пішохідна маршрутизація", "Edge Function route-hike; модуль _shared/routing.ts"],
        ["Персоналізовані рекомендації", "Edge Function recommend-routes"],
        ["Підготовка офлайн-треку", "Edge Function prepare-offline-route"],
        ["Групові походи (заявки, рішення)", "Edge Function trip-actions"],
        ["Чат походу", "Edge Function trip-chat"],
        ["Прогноз погоди", "Edge Function weather"],
        ["Геопошук місць", "Edge Function geosearch"],
        ["POI на карті (вершини)", "Edge Function poi-nearby"],
        ["ШІ-консультант", "Edge Function ai-chat"],
        ["In-app сповіщення", "таблиця notifications + insert через service role"],
        ["Ізоляція даних користувачів", "RLS на всіх таблицях"],
    ],
)

h2("4.1.1. Реалізація серверних функцій предметної області")
body(
    "Кожна Edge Function реалізує одну або кілька функцій з табл. 4.0. Нижче описано "
    "результати розробки: вхідні дані, алгоритм обробки на сервері та вихідний результат. "
    "Уривки коду ілюструють ключові обчислення; повні файли — у каталозі supabase/functions/."
)

h3("Функція збереження та оновлення маршруту")
body(
    "Реалізовано функцію створення й редагування маршруту користувачем (вимоги розд. 2 "
    "щодо каталогу та редактора). Програмний модуль — Edge Function save-route. "
    "На сервері виконується: перевірка наявності точок «старт» і «фініш»; обчислення "
    "distance_km, ascent_m, duration_h (computeRouteStats); побудова GeoJSON-треку "
    "(fetchHikingRouteThrough: GraphHopper → OSRM → сегменти); атомарний запис у таблиці "
    "routes і route_points. Клієнт передає масив points і отримує route_id."
)
code(
    """const stats = computeRouteStats(statsInput);
const { points: poly } = await fetchHikingRouteThrough(waypoints);
const geojson = toGeoJsonLineString(poly);
await supabase.from("routes").insert({
  ...routePayload, author_id: user.id, geojson,
});"""
)

h3("Функція пішохідної маршрутизації")
body(
    "Реалізовано функцію побудови лінії руху по стежках для редактора та навігації. "
    "Модуль route-hike приймає waypoints або route_id, викликає fetchHikingRouteThrough "
    "і повертає polyline та джерело (graphhopper | osrm | chain). Ланцюг резервування "
    "забезпечує роботу при недоступності окремого постачальника."
)
code(
    """export async function fetchHikingRouteThrough(waypoints) {
  if (ghKey()) try {
    return { points: await graphHopperRoute(waypoints), source: "graphhopper" };
  } catch (_) {}
  try {
    return { points: await osrmFootRoute(waypoints), source: "osrm" };
  } catch (_) {}
  return { points: await chainLegRoutes(waypoints), source: "chain" };
}"""
)

h3("Функція персоналізованих рекомендацій маршрутів")
body(
    "Реалізовано функцію підбору маршрутів на головному екрані з урахуванням профілю "
    "(fitness_level, preferred_difficulty, досвід). Модуль recommend-routes завантажує "
    "каталог публічних маршрутів і профіль за JWT (без передачі user_id у тілі); "
    "за наявності OPENAI_API_KEY формує до п’яти рекомендацій через gpt-4o-mini, "
    "інакше — rankRoutesByProfile з ваговими критеріями складності та тривалості."
)
code(
    """const recommendations = rankRoutesByProfile(profile, routeRows, 5);
return jsonResponse({ recommendations, source: "profile" });"""
)

h3("Функція підготовки офлайн-треку")
body(
    "Реалізовано функцію формування полілінії для офлайн-пакета перед завантаженням "
    "тайлів на пристрої. Модуль prepare-offline-route читає route_points, будує трек "
    "через fetchHikingRouteThrough (резерв — geojson маршруту) і повертає координати "
    "для запису в route_path.json на клієнті."
)

h3("Функції управління груповим походом і чатом")
body(
    "Реалізовано функції створення походу, подання заявки, схвалення/відхилення, "
    "закриття та скасування. Модуль trip-actions (дії create, apply, decide, cancel, "
    "close, complete, leave) перевіряє бізнес-правила (ліміт max_members, статус open), "
    "оновлює trips/trip_participants і створює записи в notifications через service role. "
    "Модуль trip-chat реалізує функції list/send повідомлень у таблиці messages з "
    "перевіркою доступу учасника зі статусом approved."
)
code(
    """await insertNotification(service, {
  user_id: trip.organizer_id,
  type: "trip_request",
  title: `Заявка від ${name}`,
  payload: { trip_id: tripId, applicant_id: userId },
});"""
)

h3("Функції інтеграції із зовнішніми сервісами")
body(
    "Реалізовано функції прогнозу погоди (weather: current, forecast, route, city), "
    "геопошуку (geosearch — Nominatim/Overpass), відображення POI (poi-nearby) та "
    "діалогового ШІ-консультанта (ai-chat). Ключі API зберігаються в Supabase Secrets; "
    "клієнт отримує лише готові JSON-відповіді через BackendApi.invoke()."
)

h2("4.1.2. Реалізація функцій доступу до даних і сповіщень")
body(
    "Реалізовано функцію ізольованого доступу до даних: на кожній таблиці увімкнено RLS "
    "(наприклад, journal_entries — лише власник; messages — учасники approved або "
    "організатор). Реалізовано функцію доставки подій у реальному часі: Supabase Realtime "
    "на таблицях trips, trip_participants, messages, notifications — клієнтські провайдери "
    "оновлюють UI без перезавантаження. Реалізовано функцію зберігання аватара: bucket "
    "avatars у Storage, метадані в profiles.avatar_url; офлайн-метадані — таблиця "
    "offline_routes (файли пакета — на пристрої)."
)

# ═══════════════════════════════════════════════════════════════
# 4.2 КЛІЄНТ — реалізовані функції
# ═══════════════════════════════════════════════════════════════
h2("4.2. Реалізація клієнтської частини системи")
body(
    "Клієнтську частину реалізовано мовою Dart у фреймворку Flutter [11]. "
    "Програмний код організовано за модулями lib/features/ (auth, routes, navigation, "
    "weather, trips, journal, profile, ai, home). У табл. 4.0а наведено відповідність "
    "функцій специфікації та екранів/класів реалізації."
)
add_table(
    "Таблиця 4.0а. Відповідність функцій специфікації та програмної реалізації (клієнт)",
    ["Функція системи", "Екран / програмний компонент"],
    [
        ["Вхід, реєстрація, OAuth", "LoginScreen, RegisterScreen; app_router redirect"],
        ["Каталог і фільтрація маршрутів", "RoutesScreen; RoutesRepository"],
        ["Деталі маршруту, офлайн, старт", "RouteDetailsScreen"],
        ["Редактор маршруту", "Route editor; SaveRouteApi → save-route"],
        ["Онлайн-навігація, reroute", "NavigationScreen; RoutingRepository"],
        ["Офлайн-завантаження карт", "OfflineMapService; prepare-offline-route"],
        ["Офлайн-навігація", "NavigationScreen (offline=true)"],
        ["Групові походи, чат", "GroupHikesScreen, TripChatScreen; TripsApi"],
        ["Рекомендації та ШІ", "HomeScreen, AiChatPanel"],
        ["Погода", "RouteWeatherScreen; WeatherRepository"],
        ["Журнал походів", "JournalScreen; journal_entries"],
        ["Профіль, статистика, досягнення", "ProfileScreen, EditProfileScreen"],
    ],
)

h2("4.2.1. Реалізація функцій автентифікації та захисту сеансу")
body(
    "Реалізовано функцію входу за email/паролем і через Google OAuth (signInWithOAuth). "
    "Реалізовано функцію примусового завершення налаштування профілю після OAuth: "
    "GoRouter.redirect перевіряє поле age у profiles і перенаправляє на /register?oauth=1. "
    "Реалізовано функцію блокування неавторизованого доступу до вкладок застосунку."
)
code(
    """redirect: (context, state) async {
  final session = Supabase.instance.client.auth.currentSession;
  if (session == null) return '/login';
  if (await _needsPhysicalProfile(session.user.id))
    return '/register?oauth=1';
  return null;
}"""
)

h2("4.2.2. Реалізація функцій роботи з маршрутами")
body(
    "Реалізовано функцію перегляду каталогу з пошуком і фільтрами (складність, тип, "
    "тривалість, набір висоти) — динамічний запит PostgREST у RoutesRepository. "
    "Реалізовано функцію геопошуку точок редактора — виклик geosearch. "
    "Реалізовано функцію побудови треку в редакторі — route-hike. "
    "Реалізовано функцію збереження маршруту — SaveRouteApi викликає save-route."
)
code(
    """final data = await _api.invoke('save-route', body: {
  'action': 'create', 'title': title, 'points': points,
}, timeout: const Duration(seconds: 120));"""
)

h2("4.2.3. Реалізація функцій офлайн-карт і навігації")
body(
    "Реалізовано функцію завантаження офлайн-пакета: prepare-offline-route → bbox → "
    "пакетне завантаження тайлів OSM (OfflineMapService), маркер .complete, upsert "
    "offline_routes. Реалізовано функцію онлайн-навігації з GPS, відображенням "
    "пройденого/залишкового треку та POI (poi-nearby). Реалізовано функцію автоматичної "
    "перебудови маршруту при відхиленні (пороги 45 м / 75 м, offline=true вимикає reroute)."
)
code(
    """static const double _offRouteThresholdM = 45;
static const double _offRouteRerouteImmediatelyM = 75;
if (_offlineOnlyNav) return;
final newRoute = await _routingRepo.fetchHikingRouteThrough([user, ...remaining]);"""
)

h2("4.2.4. Реалізація функцій групових походів")
body(
    "Реалізовано функції перегляду списку походів, подання заявки, схвалення учасників "
    "організатором, групового чату — TripsApi → trip-actions, trip-chat; оновлення списку "
    "і чату через Realtime. In-app сповіщення про заявки — SnackBar (InAppNotificationListener)."
)

h2("4.2.5. Реалізація функцій погоди, журналу та профілю")
body(
    "Реалізовано функцію перегляду погоди за координатами та для маршруту (weather). "
    "Реалізовано функцію ведення журналу: запис після навігації з метриками, до 5 фото "
    "(journal_entries, journal_photos). Реалізовано функцію профілю та агрегованої "
    "статистики (profile_stats), досягнень, вкладки «Мої маршрути» з офлайн-пакетами. "
    "Повний код NavigationScreen, RouteDetailsScreen — у репозиторії lib/features/; "
    "структура модулів — у додатку Б."
)

doc.add_page_break()


def append_v8_from(section_marker: str) -> None:
    """Текст з v8 від заданого підрозділу; таблиці — як Word-таблиці."""
    if not V8_TAIL.exists():
        body("[Скопіюйте п. 4.3–4.5 з rozd4_hikora (8).docx]")
        return

    text = V8_TAIL.read_text(encoding="utf-8")
    start = text.find(section_marker)
    if start < 0:
        return
    chunk = text[start:]
    lines = [ln.strip() for ln in chunk.splitlines()]

    headers_h2 = tuple(f"4.{i}." for i in range(3, 6)) + (
        "4.4.1.",
        "4.4.1.1.",
        "4.4.1.2.",
        "4.4.2.",
        "4.4.3.",
        "4.4.4.",
        "4.4.5.",
        "4.3.1.",
        "4.3.2.",
        "4.3.3.",
        "4.3.4.",
        "4.3.5.",
        "Додаток Л.",
        "Л.1.",
        "Л.2.",
        "Л.3.",
    )

    i = 0
    buf: list[str] = []

    def flush_para() -> None:
        nonlocal buf
        if not buf:
            return
        para = " ".join(buf).strip()
        buf = []
        if not para:
            return
        if para.startswith("Рис."):
            caption(para)
        elif any(para.startswith(h) for h in headers_h2) and len(para) < 130:
            if para.startswith("4.4.1.") or para.startswith("4.3."):
                h3(para)
            else:
                h2(para)
        else:
            body(para)

    def read_table_at(pos: int) -> int:
        """Парсер таблиць: Таблиця 4.x → заголовок → header row → пари рядків."""
        nonlocal i
        cap_lines = []
        if lines[pos].startswith("Таблиця"):
            cap_lines.append(lines[pos])
            pos += 1
        while pos < len(lines) and lines[pos] and not lines[pos].startswith("Таблиця"):
            if any(lines[pos].startswith(f"4.{n}") for n in range(3, 6)):
                break
            if lines[pos].startswith("4.4.") and not lines[pos][:6] in (
                "4.4.1.",
                "4.4.2",
            ):
                break
            cap_lines.append(lines[pos])
            pos += 1
        if len(cap_lines) < 3:
            return pos
        cap = " ".join(cap_lines[:2]) if len(cap_lines) >= 2 else cap_lines[0]
        rest = [x for x in cap_lines[2:] if x]
        if not rest:
            return pos
        # 2-3 колонки: перший рядок після підпису — заголовки
        if len(rest) >= 2 and len(rest) % 2 == 0 and len(rest[0]) < 40:
            hdr = [rest[0], rest[1]]
            data_rows = []
            for k in range(2, len(rest), 2):
                data_rows.append([rest[k], rest[k + 1]])
            add_table(cap, hdr, data_rows)
        elif len(rest) >= 3 and len(rest) % 3 == 0:
            hdr = [rest[0], rest[1], rest[2]]
            data_rows = []
            for k in range(3, len(rest), 3):
                data_rows.append([rest[k], rest[k + 1], rest[k + 2]])
            add_table(cap, hdr, data_rows)
        else:
            for ln in cap_lines:
                caption(ln)
        return pos

    while i < len(lines):
        s = lines[i]
        if not s:
            flush_para()
            i += 1
            continue
        if s.startswith("Таблиця"):
            flush_para()
            i = read_table_at(i)
            continue
        if re.match(r"^function |^export |^const |^await |^if \(|^for \(", s):
            flush_para()
            block = []
            while i < len(lines) and lines[i] and (
                re.match(r"^(function |export |const |await |if |for |return |let |})",
                         lines[i])
                or lines[i].startswith("    ")
                or ";" in lines[i]
                or lines[i].endswith("}") or lines[i].endswith("},")
            ):
                block.append(lines[i])
                i += 1
            if block:
                code("\n".join(block[:35]))
            continue
        if buf and any(s.startswith(h) for h in headers_h2) and len(s) < 120:
            flush_para()
        buf.append(s)
        i += 1
    flush_para()


# 4.3 — демонстрація; 4.4 — тестування за структурою додатка Л
append_v8_from("4.3. Результати реалізації")

# Підсилення вступу до 4.4 за п. 3.12.3 (якщо в v8 вже є — дубль не шкодить короткому абзацу)
if V8_TAIL.read_text(encoding="utf-8").find("4.4. Результати тестування") >= 0:
    pass  # уже в append

doc.save(OUT)
print(f"Saved: {OUT}")
